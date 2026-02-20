#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
grammar.py - Document quality auditor for DOCX/PDF/TXT/MD.
Generates an issues report and a 0-100 score based on grammar, style,
readability, and consistency checks.

USAGE:
    python grammar.py /path/to/document.docx --lang en-US --out /path/to/report_prefix
    python grammar.py --config /path/to/config.toml
    python grammar.py --enable language_tool --html --sarif

OUTPUTS:
    - <prefix>_report.json   (machine-readable findings and scores)
    - <prefix>_report.md     (human-readable summary)
    - <prefix>_report.html   (optional HTML report)
    - <prefix>_report.sarif  (optional SARIF report)
    - prints a compact summary to stdout

OPTIONAL DEPENDENCIES (install what you need):
    pip install python-docx pdfminer.six PyPDF2 language-tool-python textstat
    pip install pytesseract pdf2image  # OCR (requires tesseract + poppler)
    pip install pyyaml  # YAML config support

NOTES:
    - If language_tool_python is installed, a local LanguageTool server will be started.
      For best performance, ensure Java is installed.
    - If both pdfminer and PyPDF2 are available, pdfminer is preferred for text extraction.
    - OCR is optional; use --ocr auto|force for scanned PDFs.
"""

from __future__ import annotations

import argparse
import bisect
import html
import json
import logging
import os
import re
import statistics
import sys
import unicodedata
from dataclasses import dataclass, asdict
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

# Optional imports
try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
except Exception:
    pdfminer_extract_text = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import language_tool_python  # grammar/spelling/style (rules-based)
except Exception:
    language_tool_python = None

try:
    import textstat  # readability
except Exception:
    textstat = None

try:
    import pytesseract  # OCR
except Exception:
    pytesseract = None

try:
    from pdf2image import convert_from_path  # OCR PDF rendering
except Exception:
    convert_from_path = None

try:
    import tomllib  # Python 3.11+
except Exception:
    tomllib = None

try:
    import yaml  # PyYAML (optional)
except Exception:
    yaml = None

try:
    from pdfminer.pdfpage import PDFPage
except Exception:
    PDFPage = None

LOGGER = logging.getLogger("doc_audit")

EM_DASH = "\u2014"
ELLIPSIS = "\u2026"
DEFAULT_FILE = "document.pdf"

ABBREVIATIONS = [
    "e.g.", "i.e.", "mr.", "mrs.", "ms.", "dr.", "prof.", "sr.", "jr.",
    "vs.", "no.", "u.s.", "u.k.", "etc.", "inc.", "ltd.", "jan.", "feb.",
    "mar.", "apr.", "jun.", "jul.", "aug.", "sep.", "sept.", "oct.", "nov.", "dec."
]
ABBREV_RE = re.compile(r"\b(?:" + "|".join(re.escape(a) for a in ABBREVIATIONS) + r")", re.IGNORECASE)
WORD_RE = re.compile(r"[^\W\d_]+(?:['-][^\W\d_]+)*", flags=re.UNICODE)


# -----------------------------
# Data structures
# -----------------------------

@dataclass
class Issue:
    check_id: str
    message: str
    location: Optional[str] = None  # e.g., "Line 42, Col 5" or a short snippet
    severity: Optional[str] = None
    line: Optional[int] = None
    col: Optional[int] = None
    page: Optional[int] = None
    snippet: Optional[str] = None
    highlight: Optional[str] = None

@dataclass
class CheckResult:
    check_id: str
    name: str
    score: float
    max_score: float
    issues: List[Issue]
    metrics: Dict[str, Any]
    status: str  # ok | skipped | error

@dataclass
class AuditReport:
    file_path: str
    language: str
    total_score: float
    coverage: float
    check_results: List[CheckResult]
    summary: Dict[str, Any]
    doc_info: Dict[str, Any]

@dataclass
class CheckSpec:
    check_id: str
    name: str
    weight: float
    func: Any

@dataclass
class AuditOptions:
    max_issues: int
    max_issues_per_check: int
    lt_chunk_size: int
    lt_overlap: int
    strict: bool
    enable_langtool: bool
    ocr_mode: str
    ocr_lang: str
    ocr_dpi: int
    ocr_min_chars: int

@dataclass
class CheckContext:
    text: str
    language: str
    lines: List[str]
    sentences: List[str]
    words: List[str]
    line_starts: List[int]
    page_breaks: List[int]
    options: AuditOptions


@dataclass
class DocumentData:
    text: str
    page_count: Optional[int] = None
    ocr_used: bool = False
    ocr_reason: Optional[str] = None
    extraction_method: Optional[str] = None


# -----------------------------
# Helpers
# -----------------------------

def read_text_from_file(file_path: str, options: AuditOptions) -> DocumentData:
    ext = os.path.splitext(file_path)[1].lower()
    if ext in [".txt", ".md"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return DocumentData(text=f.read(), extraction_method="text")
    if ext == ".docx":
        if docx is None:
            raise RuntimeError("python-docx is required to read .docx files. Install with: pip install python-docx")
        return DocumentData(text=_read_docx(file_path), extraction_method="docx")
    if ext == ".pdf":
        return _read_pdf(file_path, options)
    # Attempt to read as text fallback
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return DocumentData(text=f.read(), extraction_method="text")


def _read_docx(file_path: str) -> str:
    d = docx.Document(file_path)
    parts = []
    for p in d.paragraphs:
        if p.text:
            parts.append(p.text)
    # Include simple table cell extraction
    for t in d.tables:
        for row in t.rows:
            cells = [c.text for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_pdf(file_path: str, options: AuditOptions) -> DocumentData:
    page_texts: List[str] = []
    extraction_method = None
    page_count = None

    # Prefer pdfminer when available, but keep page boundaries.
    if pdfminer_extract_text is not None and PDFPage is not None:
        try:
            with open(file_path, "rb") as f:
                pages = list(PDFPage.get_pages(f))
            page_count = len(pages)
            for i in range(page_count):
                page_texts.append(pdfminer_extract_text(file_path, page_numbers=[i]) or "")
            extraction_method = "pdfminer"
        except Exception:
            page_texts = []

    if not page_texts:
        if PyPDF2 is None:
            raise RuntimeError(
                "Reading PDFs requires either pdfminer.six or PyPDF2. Install with: pip install pdfminer.six or pip install PyPDF2"
            )
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)
            for page in reader.pages:
                try:
                    page_texts.append(page.extract_text() or "")
                except Exception:
                    page_texts.append("")
        extraction_method = "pypdf2"

    text = "\f".join(page_texts)
    doc = DocumentData(text=text, page_count=page_count, extraction_method=extraction_method)

    if options.ocr_mode != "off":
        stripped_len = len(text.strip())
        needs_ocr = options.ocr_mode == "force" or stripped_len < options.ocr_min_chars
        if needs_ocr:
            try:
                ocr_pages = ocr_pdf_to_pages(file_path, options.ocr_lang, options.ocr_dpi)
                doc.text = "\f".join(ocr_pages)
                doc.ocr_used = True
                doc.ocr_reason = "forced" if options.ocr_mode == "force" else f"low text ({stripped_len} chars)"
                doc.page_count = len(ocr_pages)
                doc.extraction_method = "ocr"
            except Exception as exc:
                LOGGER.warning("OCR failed: %s", exc)
    return doc


def ocr_pdf_to_pages(file_path: str, lang: str, dpi: int) -> List[str]:
    if pytesseract is None or convert_from_path is None:
        raise RuntimeError(
            "OCR requires pytesseract and pdf2image (plus poppler). Install with: "
            "pip install pytesseract pdf2image and ensure `tesseract`/`pdftoppm` are available."
        )
    images = convert_from_path(file_path, dpi=dpi)
    texts: List[str] = []
    for image in images:
        texts.append(pytesseract.image_to_string(image, lang=lang))
    return texts


def normalize_text(text: str) -> str:
    # Normalize unicode, standardize line breaks, remove trailing spaces
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+\n", "\n", text)
    return text


def build_line_starts(text: str) -> List[int]:
    starts = [0]
    for match in re.finditer(r"[\n\f]", text):
        starts.append(match.end())
    return starts


def offset_to_line_col(offset: int, line_starts: Sequence[int]) -> Tuple[int, int]:
    # Binary search for line start just before offset
    lo, hi = 0, len(line_starts)
    while lo < hi:
        mid = (lo + hi) // 2
        if line_starts[mid] <= offset:
            lo = mid + 1
        else:
            hi = mid
    line_idx = max(0, lo - 1)
    line_no = line_idx + 1
    col_no = offset - line_starts[line_idx] + 1
    return line_no, col_no


def offset_to_page(offset: int, page_breaks: Sequence[int]) -> int:
    if not page_breaks:
        return 1
    return bisect.bisect_right(page_breaks, offset) + 1


def offset_to_page_line_col(
    offset: int, line_starts: Sequence[int], page_breaks: Sequence[int]
) -> Tuple[Optional[int], int, int]:
    line_no, col_no = offset_to_line_col(offset, line_starts)
    if not page_breaks:
        return None, line_no, col_no
    page_idx = bisect.bisect_right(page_breaks, offset)
    page_start = 0 if page_idx == 0 else page_breaks[page_idx - 1] + 1
    page_line_start_idx = max(0, bisect.bisect_right(line_starts, page_start) - 1)
    line_idx = max(0, bisect.bisect_right(line_starts, offset) - 1)
    page_line_no = max(1, line_idx - page_line_start_idx + 1)
    return page_idx + 1, page_line_no, col_no


def word_tokens(text: str) -> List[str]:
    return WORD_RE.findall(text)


def percent_values(text: str) -> List[str]:
    return re.findall(
        r"\b\d+(?:\.\d+)?\s*%|\b\d+(?:\.\d+)?\s+(?:percent|per\s*cent)\b",
        text,
        flags=re.IGNORECASE,
    )


def _protect_abbreviations(text: str) -> str:
    return ABBREV_RE.sub(lambda m: m.group(0).replace(".", "<DOT>"), text)


def split_sentences(text: str) -> List[str]:
    # Heuristic sentence splitter. Keeps common abbreviations, initials, and decimals intact.
    collapsed = re.sub(r"\s+", " ", text)
    collapsed = _protect_abbreviations(collapsed)
    collapsed = re.sub(r"(\d)\.(\d)", r"\1<DECIMAL>\2", collapsed)
    collapsed = re.sub(r"\b([A-Z])\.", r"\1<DOT>", collapsed)
    boundary_re = re.compile(r"(?<=[.!?])\s+(?=[\"'\(\[]?[A-Z0-9])")
    candidates = boundary_re.split(collapsed)
    results = []
    for c in candidates:
        c = c.replace("<DOT>", ".").replace("<DECIMAL>", ".").strip()
        if c and not c.isspace():
            results.append(c)
    return results


def iter_text_chunks(text: str, max_chars: int, overlap: int) -> Iterable[Tuple[str, int]]:
    if max_chars <= 0 or len(text) <= max_chars:
        yield text, 0
        return
    start = 0
    length = len(text)
    while start < length:
        end = min(length, start + max_chars)
        if end < length:
            # Try to cut on a boundary to reduce grammar-tool errors
            boundary = text.rfind("\n", start, end)
            if boundary == -1:
                boundary = text.rfind(" ", start, end)
            if boundary != -1 and boundary > start + max_chars // 2:
                end = boundary
        chunk = text[start:end]
        yield chunk, start
        if end == length:
            break
        start = max(0, end - overlap)


def format_location(line: int, col: int, snippet: Optional[str] = None, page: Optional[int] = None) -> str:
    if page is not None:
        base = f"Page {page}, Line {line}, Col {col}"
    else:
        base = f"Line {line}, Col {col}"
    if snippet:
        return f"{base}: {snippet}"
    return base


def snippet_around(text: str, start: int, end: int, width: int = 80) -> str:
    half = max(10, width // 2)
    left = max(0, start - half)
    right = min(len(text), end + half)
    snippet = text[left:right].replace("\n", " ").replace("\f", " ").strip()
    return snippet


def issue_from_span(
    ctx: CheckContext,
    check_id: str,
    message: str,
    start: int,
    end: int,
    width: int = 80,
) -> Issue:
    snippet = snippet_around(ctx.text, start, end, width)
    highlight = ctx.text[start:end].replace("\n", " ").replace("\f", " ").strip()
    page_no, line_no, col_no = offset_to_page_line_col(start, ctx.line_starts, ctx.page_breaks)
    location = format_location(line_no, col_no, snippet if snippet else None, page=page_no)
    return Issue(
        check_id,
        message,
        location=location,
        line=line_no,
        col=col_no,
        page=page_no,
        snippet=snippet if snippet else None,
        highlight=highlight if highlight else None,
    )


def add_match_issues(
    ctx: CheckContext,
    issues: List[Issue],
    check_id: str,
    pattern: str,
    message: Any,
    max_samples: int = 3,
    flags: int = 0,
) -> int:
    added = 0
    for m in re.finditer(pattern, ctx.text, flags):
        if len(issues) >= ctx.options.max_issues_per_check:
            break
        msg = message(m) if callable(message) else message
        issues.append(issue_from_span(ctx, check_id, msg, m.start(), m.end()))
        added += 1
        if added >= max_samples:
            break
    return added


def langtool_error_hint(exc: Exception) -> Optional[str]:
    msg = str(exc).lower()
    if "java" in msg or "jre" in msg or "jvm" in msg:
        return "Install a JRE/JDK and ensure `java` is on PATH."
    if "connection" in msg or "refused" in msg or "port" in msg:
        return "LanguageTool server failed to start; check Java install and localhost access."
    if "timeout" in msg:
        return "LanguageTool timed out; try smaller --lt-chunk-size."
    return None


def build_context(text: str, language: str, options: AuditOptions) -> CheckContext:
    lines = re.split(r"[\n\f]", text)
    sentences = split_sentences(text)
    words = word_tokens(text)
    line_starts = build_line_starts(text)
    page_breaks = [m.start() for m in re.finditer(r"\f", text)]
    return CheckContext(
        text=text,
        language=language,
        lines=lines,
        sentences=sentences,
        words=words,
        line_starts=line_starts,
        page_breaks=page_breaks,
        options=options,
    )


# -----------------------------
# Checks
# -----------------------------

def check_language_tool(ctx: CheckContext, max_points: float) -> CheckResult:
    issues: List[Issue] = []
    metrics: Dict[str, Any] = {"matches": 0}

    if not ctx.options.enable_langtool or language_tool_python is None:
        return CheckResult(
            check_id="language_tool",
            name="Grammar and style (LanguageTool)",
            score=0.0,
            max_score=max_points,
            issues=[
                Issue(
                    "language_tool",
                    "language_tool_python not installed or disabled; skipping deep grammar/style checks.",
                )
            ],
            metrics={
                "available": False,
                "scoring_notes": "Not scored because LanguageTool is unavailable or disabled.",
            },
            status="skipped",
        )

    tool = None
    try:
        tool = language_tool_python.LanguageTool(ctx.language)
        words = max(1, len(ctx.words))
        total_matches = 0
        seen: set = set()

        for chunk, base_offset in iter_text_chunks(ctx.text, ctx.options.lt_chunk_size, ctx.options.lt_overlap):
            matches = tool.check(chunk)
            for m in matches:
                global_offset = base_offset + m.offset
                key = (m.ruleId, global_offset, getattr(m, "errorLength", None))
                if key in seen:
                    continue
                seen.add(key)
                total_matches += 1
                if len(issues) >= ctx.options.max_issues_per_check:
                    continue
                msg = f"{m.ruleId}: {m.message}"
                if m.replacements:
                    msg += f"; suggestions: {', '.join(m.replacements[:3])}"
                error_len = getattr(m, "errorLength", 1) or 1
                issues.append(
                    issue_from_span(
                        ctx,
                        "language_tool",
                        msg,
                        global_offset,
                        global_offset + error_len,
                    )
                )

        metrics["matches"] = total_matches
        matches_per_1k = 1000 * total_matches / words
        metrics["matches_per_1000_words"] = round(matches_per_1k, 2)

        penalty = min(max_points, matches_per_1k / 2.0)
        metrics["penalty_points"] = round(penalty, 2)
        metrics["scoring_notes"] = "Penalty = matches_per_1000_words / 2 (capped at max score)."
        score = max_points - penalty
        return CheckResult(
            check_id="language_tool",
            name="Grammar and style (LanguageTool)",
            score=round(score, 2),
            max_score=max_points,
            issues=issues,
            metrics=metrics,
            status="ok",
        )
    except Exception as exc:
        hint = langtool_error_hint(exc)
        issues = [Issue("language_tool", f"LanguageTool error: {exc}")]
        if hint:
            issues.append(Issue("language_tool", f"Suggested fix: {hint}"))
        return CheckResult(
            check_id="language_tool",
            name="Grammar and style (LanguageTool)",
            score=0.0,
            max_score=max_points,
            issues=issues,
            metrics={
                "available": False,
                "error": str(exc),
                "suggested_fix": hint,
                "scoring_notes": "Not scored because LanguageTool errored.",
            },
            status="error",
        )
    finally:
        if tool is not None:
            try:
                tool.close()
            except Exception:
                pass


def check_readability(ctx: CheckContext, max_points: float) -> CheckResult:
    issues: List[Issue] = []
    metrics: Dict[str, Any] = {}
    sents = ctx.sentences
    words = ctx.words
    avg_len = statistics.mean([len(word_tokens(s)) for s in sents]) if sents else 0
    long_sents = [s for s in sents if len(word_tokens(s)) > 30]
    metrics["avg_sentence_length_words"] = round(avg_len, 2)
    metrics["long_sentences_over_30_words"] = len(long_sents)
    metrics["word_count"] = len(words)
    metrics["sentence_count"] = len(sents)

    if textstat is not None:
        try:
            fre = textstat.flesch_reading_ease(ctx.text)
            grade = textstat.text_standard(ctx.text, float_output=True)
            metrics["flesch_reading_ease"] = round(fre, 2)
            metrics["grade_level"] = round(grade, 2)
        except Exception:
            pass

    penalty = min(max_points, 0.5 * len(long_sents))
    metrics["penalty_points"] = round(penalty, 2)
    metrics["scoring_notes"] = "0.5 points per sentence over 30 words (capped)."
    score = max_points - penalty

    for ex in long_sents[: min(3, ctx.options.max_issues_per_check)]:
        issues.append(
            Issue(
                "readability",
                "Very long sentence (over 30 words). Consider splitting.",
                location=ex[:120] + ("..." if len(ex) > 120 else ""),
                snippet=ex[:200] + ("..." if len(ex) > 200 else ""),
            )
        )

    return CheckResult(
        check_id="readability",
        name="Readability",
        score=round(score, 2),
        max_score=max_points,
        issues=issues,
        metrics=metrics,
        status="ok",
    )


def check_punctuation_style(ctx: CheckContext, max_points: float) -> CheckResult:
    issues: List[Issue] = []
    metrics: Dict[str, Any] = {}
    words = max(1, len(ctx.words))
    semicolons = ctx.text.count(";")
    ellipses = ctx.text.count("...") + ctx.text.count(ELLIPSIS)
    double_punct = len(re.findall(r"[!?]{2,}", ctx.text))
    em_dashes = ctx.text.count(EM_DASH)
    double_hyphens = len(re.findall(r"--", ctx.text))
    spaced_em_dashes = len(re.findall(rf"\s{EM_DASH}\s", ctx.text))
    spaced_double_hyphens = len(re.findall(r"\s--\s", ctx.text))

    metrics.update({
        "semicolons": semicolons,
        "semicolons_per_1000_words": round(1000 * semicolons / words, 2),
        "ellipses": ellipses,
        "double_punctuations": double_punct,
        "em_dashes": em_dashes,
        "double_hyphens": double_hyphens,
        "em_dashes_with_spaces": spaced_em_dashes,
        "double_hyphens_with_spaces": spaced_double_hyphens,
    })

    penalty = 0.0
    penalty_reasons: List[str] = []
    if semicolons > 0 and (1000 * semicolons / words) > 3.0:
        penalty += 2.5
        penalty_reasons.append("Semicolon density > 3 per 1,000 words (-2.5).")
        add_match_issues(
            ctx,
            issues,
            "punctuation",
            r";",
            "Semicolon usage (sample contributing to density penalty).",
        )
    if ellipses > 2:
        penalty += 1.5
        penalty_reasons.append("Ellipses appear frequently (-1.5).")
        add_match_issues(
            ctx,
            issues,
            "punctuation",
            r"\.\.\.|…",
            "Ellipsis usage (sample).",
        )
    if double_punct > 0:
        penalty += 1.0
        penalty_reasons.append("Repeated punctuation (e.g., !! or ??) (-1.0).")
        add_match_issues(
            ctx,
            issues,
            "punctuation",
            r"[!?]{2,}",
            "Repeated punctuation (sample).",
        )
    if spaced_em_dashes > 0 or spaced_double_hyphens > 0:
        penalty += 0.5
        penalty_reasons.append("Spaced em dashes or double hyphens found (-0.5).")
        add_match_issues(
            ctx,
            issues,
            "punctuation",
            rf"\s{EM_DASH}\s|\s--\s",
            "Spaced dash usage (sample).",
        )

    applied_penalty = min(max_points, penalty)
    metrics["penalty_points"] = round(applied_penalty, 2)
    if penalty_reasons:
        metrics["penalty_reasons"] = penalty_reasons
    metrics["scoring_notes"] = "Penalties applied per punctuation rule (capped at max score)."
    score = max_points - applied_penalty
    return CheckResult(
        check_id="punctuation",
        name="Punctuation and style",
        score=round(score, 2),
        max_score=max_points,
        issues=issues,
        metrics=metrics,
        status="ok",
    )


def check_whitespace_formatting(ctx: CheckContext, max_points: float) -> CheckResult:
    issues: List[Issue] = []
    metrics: Dict[str, Any] = {}
    trailing_matches = [i + 1 for i, line in enumerate(ctx.lines) if re.search(r"[ \t]+$", line)]
    doubles = len(re.findall(r" {2,}", ctx.text))
    tabs = ctx.text.count("\t")
    metrics.update({
        "trailing_whitespace_lines": len(trailing_matches),
        "double_spaces": doubles,
        "tab_characters": tabs,
    })
    penalty = 0.1 * len(trailing_matches) + 0.05 * doubles + 0.1 * tabs
    applied_penalty = min(max_points, penalty)
    metrics["penalty_points"] = round(applied_penalty, 2)
    metrics["scoring_notes"] = "0.1 per trailing whitespace line, 0.05 per double space, 0.1 per tab."
    score = max_points - applied_penalty

    for line_no in trailing_matches[: ctx.options.max_issues_per_check]:
        line = ctx.lines[line_no - 1] if 0 <= line_no - 1 < len(ctx.lines) else ""
        col = len(line) if line else 1
        offset = ctx.line_starts[line_no - 1] if 0 <= line_no - 1 < len(ctx.line_starts) else 0
        page_no = offset_to_page(offset, ctx.page_breaks) if ctx.page_breaks else None
        issues.append(
            Issue(
                "whitespace",
                "Trailing whitespace detected.",
                location=format_location(line_no, col, line.strip()[:80] or None, page=page_no),
                line=line_no,
                col=col,
                page=page_no,
                snippet=line.strip()[:200] or None,
            )
        )
    if doubles > 0:
        issues.append(Issue("whitespace", f"{doubles} double-space occurrence(s) detected. Consider single spacing."))
        for line_no, line in enumerate(ctx.lines, 1):
            if len(issues) >= ctx.options.max_issues_per_check:
                break
            if "  " in line:
                col = line.find("  ") + 1
                snippet = line.strip()[:80]
                offset = ctx.line_starts[line_no - 1] + col - 1 if 0 <= line_no - 1 < len(ctx.line_starts) else 0
                page_no = offset_to_page(offset, ctx.page_breaks) if ctx.page_breaks else None
                issues.append(
                    Issue(
                        "whitespace",
                        "Double space detected (sample).",
                        location=format_location(line_no, col, snippet or None, page=page_no),
                        line=line_no,
                        col=col,
                        page=page_no,
                        snippet=snippet or None,
                        highlight="  ",
                    )
                )
    if tabs > 0:
        issues.append(Issue("whitespace", f"{tabs} tab character(s) found. Use spaces for alignment."))
        for line_no, line in enumerate(ctx.lines, 1):
            if len(issues) >= ctx.options.max_issues_per_check:
                break
            if "\t" in line:
                col = line.find("\t") + 1
                snippet = line.strip()[:80]
                offset = ctx.line_starts[line_no - 1] + col - 1 if 0 <= line_no - 1 < len(ctx.line_starts) else 0
                page_no = offset_to_page(offset, ctx.page_breaks) if ctx.page_breaks else None
                issues.append(
                    Issue(
                        "whitespace",
                        "Tab character detected (sample).",
                        location=format_location(line_no, col, snippet or None, page=page_no),
                        line=line_no,
                        col=col,
                        page=page_no,
                        snippet=snippet or None,
                        highlight="\\t",
                    )
                )

    return CheckResult(
        check_id="whitespace",
        name="Whitespace and formatting",
        score=round(score, 2),
        max_score=max_points,
        issues=issues,
        metrics=metrics,
        status="ok",
    )


def check_percentage_consistency(ctx: CheckContext, max_points: float) -> CheckResult:
    issues: List[Issue] = []
    metrics: Dict[str, Any] = {}
    vals = percent_values(ctx.text)
    pct = len([v for v in vals if "%" in v])
    wordpct = len([v for v in vals if re.search(r"\bpercent\b|per\s*cent", v, flags=re.IGNORECASE)])
    metrics.update({
        "percent_symbol": pct,
        "percent_word": wordpct,
        "total_percent_mentions": len(vals),
    })
    penalty = 0.0
    penalty_reasons: List[str] = []
    if pct > 0 and wordpct > 0:
        penalty = 2.0
        penalty_reasons.append("Mix of '%' and 'percent' styles (-2.0).")
        add_match_issues(
            ctx,
            issues,
            "percentages",
            r"\b\d+(?:\.\d+)?\s*%",
            lambda m: f"Percent symbol style used (sample): {m.group(0)}",
        )
        add_match_issues(
            ctx,
            issues,
            "percentages",
            r"\b\d+(?:\.\d+)?\s+(?:percent|per\s*cent)\b",
            lambda m: f"'Percent' word style used (sample): {m.group(0)}",
            flags=re.IGNORECASE,
        )

    decimals = re.findall(r"\b\d+\.\d+\s*%", ctx.text)
    integers = re.findall(r"\b\d+\s*%", ctx.text)
    if decimals and integers:
        penalty += 0.5
        penalty_reasons.append("Mix of whole-number and decimal percentages (-0.5).")
        add_match_issues(
            ctx,
            issues,
            "percentages",
            r"\b\d+\.\d+\s*%",
            lambda m: f"Decimal percentage used (sample): {m.group(0)}",
        )
        add_match_issues(
            ctx,
            issues,
            "percentages",
            r"\b\d+\s*%",
            lambda m: f"Whole-number percentage used (sample): {m.group(0)}",
        )

    applied_penalty = min(max_points, penalty)
    metrics["penalty_points"] = round(applied_penalty, 2)
    if penalty_reasons:
        metrics["penalty_reasons"] = penalty_reasons
    metrics["scoring_notes"] = "Penalties applied for inconsistent percentage styles."
    score = max_points - applied_penalty
    return CheckResult(
        check_id="percentages",
        name="Numbers and percentages consistency",
        score=round(score, 2),
        max_score=max_points,
        issues=issues,
        metrics=metrics,
        status="ok",
    )


def check_headings_capitalization(ctx: CheckContext, max_points: float) -> CheckResult:
    issues: List[Issue] = []
    metrics: Dict[str, Any] = {}
    lines = ctx.lines
    headings: List[Tuple[int, str]] = []

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^#{1,6}\s+\S", stripped):
            heading = re.sub(r"^#{1,6}\s+", "", stripped)
            headings.append((i + 1, heading))
            continue
        if stripped.endswith(":") and len(stripped) < 80:
            headings.append((i + 1, stripped[:-1].strip()))
            continue
        if stripped.isupper() and len(stripped) < 60:
            headings.append((i + 1, stripped))
            continue
        if i + 1 < len(lines):
            underline = lines[i + 1].strip()
            if re.match(r"^[=-]{3,}$", underline):
                headings.append((i + 1, stripped))

    metrics["probable_headings"] = len(headings)

    bad = 0
    for line_no, heading in headings:
        if heading and heading[0].islower():
            bad += 1
            offset = ctx.line_starts[line_no - 1] if 0 <= line_no - 1 < len(ctx.line_starts) else 0
            page_no = offset_to_page(offset, ctx.page_breaks) if ctx.page_breaks else None
            issues.append(
                Issue(
                    "headings",
                    "Heading may not be capitalized.",
                    location=format_location(line_no, 1, heading[:80], page=page_no),
                    line=line_no,
                    col=1,
                    page=page_no,
                    snippet=heading[:200],
                )
            )

    penalty = min(max_points, bad * 0.5)
    metrics["penalty_points"] = round(penalty, 2)
    metrics["scoring_notes"] = "0.5 points per heading starting with lowercase (capped)."
    score = max_points - penalty
    return CheckResult(
        check_id="headings",
        name="Headings and capitalization",
        score=round(score, 2),
        max_score=max_points,
        issues=issues,
        metrics=metrics,
        status="ok",
    )


def check_passive_voice(ctx: CheckContext, max_points: float) -> CheckResult:
    issues: List[Issue] = []
    metrics: Dict[str, Any] = {}
    be_forms = r"\b(?:am|is|are|was|were|be|been|being)\b"
    participle = r"\b\w+(?:ed|en)\b"
    matches = list(re.finditer(be_forms + r"\s+" + participle, ctx.text, flags=re.IGNORECASE))
    metrics["passive_like_phrases"] = len(matches)

    penalty = min(max_points, len(matches) * 0.05)
    metrics["penalty_points"] = round(penalty, 2)
    metrics["scoring_notes"] = "0.05 points per passive-like phrase (capped)."
    score = max_points - penalty

    for m in matches[: ctx.options.max_issues_per_check]:
        phrase = m.group(0)
        issues.append(
            issue_from_span(
                ctx,
                "passive_voice",
                f"Possible passive construction: '{phrase}'",
                m.start(),
                m.end(),
            )
        )

    return CheckResult(
        check_id="passive_voice",
        name="Passive voice (heuristic)",
        score=round(score, 2),
        max_score=max_points,
        issues=issues,
        metrics=metrics,
        status="ok",
    )


def check_parallel_bullets(ctx: CheckContext, max_points: float) -> CheckResult:
    """
    Check bullet/numbered lists for parallel starts (all gerunds, all imperatives, or all noun phrases).
    Heuristic: categorize each item start into {GERUND, IMPERATIVE, NOUNPHRASE, OTHER}.
    """
    issues: List[Issue] = []
    metrics: Dict[str, Any] = {}

    lines = ctx.lines
    bullet_blocks: List[List[Tuple[int, str]]] = []
    block: List[Tuple[int, str]] = []
    bullet_re = re.compile(r"^\s*(?:[-*+\u2022]|\d+[.)])\s+")

    def flush_block() -> None:
        nonlocal block
        if len(block) >= 3:
            bullet_blocks.append(block)
        block = []

    for idx, ln in enumerate(lines, 1):
        if bullet_re.match(ln):
            block.append((idx, ln.strip()))
        else:
            flush_block()
    flush_block()

    def classify_start(token: str) -> str:
        t = token.lower()
        if t.endswith("ing"):
            return "GERUND"
        common_base_verbs = {
            "use", "add", "ensure", "consider", "avoid", "include", "provide",
            "attend", "highlight", "outline", "emphasize", "leverage", "budget",
            "remember", "keep", "create", "define", "verify", "review",
        }
        if t in common_base_verbs:
            return "IMPERATIVE"
        determiners = {"a", "an", "the", "this", "that", "these", "those"}
        if t in determiners:
            return "NOUNPHRASE"
        return "OTHER"

    inconsistent_blocks = 0
    for blk in bullet_blocks:
        starts = []
        for _, item in blk:
            first = word_tokens(item[:60])
            token = first[0] if first else ""
            starts.append(classify_start(token))
        categories = set(starts)
        if len(categories) > 1:
            inconsistent_blocks += 1
            sample_lines = ", ".join(str(line_no) for line_no, _ in blk[:3])
            sample_items = " | ".join(item for _, item in blk[:3])
            first_line_no = blk[0][0]
            offset = ctx.line_starts[first_line_no - 1] if 0 <= first_line_no - 1 < len(ctx.line_starts) else 0
            page_no = offset_to_page(offset, ctx.page_breaks) if ctx.page_breaks else None
            issues.append(
                Issue(
                    "parallel_lists",
                    f"Bullet list with mixed starts ({', '.join(sorted(categories))}). Consider making items parallel.",
                    location=format_location(
                        first_line_no,
                        1,
                        f"Lines {sample_lines}: {sample_items}" + (" ..." if len(blk) > 3 else ""),
                        page=page_no,
                    ),
                    line=first_line_no,
                    col=1,
                    page=page_no,
                    snippet=sample_items + (" ..." if len(blk) > 3 else ""),
                )
            )

    metrics["bullet_blocks_3plus"] = len(bullet_blocks)
    metrics["inconsistent_blocks"] = inconsistent_blocks

    penalty = min(max_points, inconsistent_blocks * 2.0)
    metrics["penalty_points"] = round(penalty, 2)
    metrics["scoring_notes"] = "2 points per inconsistent bullet block (capped)."
    score = max_points - penalty
    return CheckResult(
        check_id="parallel_lists",
        name="Parallel structure in lists",
        score=round(score, 2),
        max_score=max_points,
        issues=issues,
        metrics=metrics,
        status="ok",
    )


def check_links_formatting(ctx: CheckContext, max_points: float) -> CheckResult:
    issues: List[Issue] = []
    metrics: Dict[str, Any] = {}
    raw_urls = re.findall(r"https?://\S+|www\.\S+", ctx.text)
    markdown_links = re.findall(r"\[[^\]]+\]\(\s*https?://[^\)]+\)", ctx.text)
    metrics["raw_urls"] = len(raw_urls)
    metrics["markdown_or_hyperlinks"] = len(markdown_links)

    penalty = 0.0
    penalty_reasons: List[str] = []
    if raw_urls and not markdown_links:
        penalty = 2.0
        penalty_reasons.append("Raw URLs shown instead of hyperlinks (-2.0).")
        add_match_issues(
            ctx,
            issues,
            "links",
            r"https?://\S+|www\.\S+",
            lambda m: f"Raw URL found (sample): {m.group(0)}",
        )

    applied_penalty = min(max_points, penalty)
    metrics["penalty_points"] = round(applied_penalty, 2)
    if penalty_reasons:
        metrics["penalty_reasons"] = penalty_reasons
    metrics["scoring_notes"] = "Penalty applies when raw URLs appear without hyperlinks."
    score = max_points - applied_penalty
    return CheckResult(
        check_id="links",
        name="Links and citation formatting",
        score=round(score, 2),
        max_score=max_points,
        issues=issues,
        metrics=metrics,
        status="ok",
    )


def _acronym_defined(text: str, acronym: str) -> bool:
    pattern1 = rf"\b[A-Za-z][A-Za-z&/ \-]{{2,80}}\s+\({re.escape(acronym)}\)"
    pattern2 = rf"\b{re.escape(acronym)}\s+\([A-Za-z][^\)]+\)"
    return re.search(pattern1, text) is not None or re.search(pattern2, text) is not None


def check_acronym_definitions(ctx: CheckContext, max_points: float) -> CheckResult:
    issues: List[Issue] = []
    metrics: Dict[str, Any] = {}
    acronyms = set(re.findall(r"\b([A-Z]{2,6})\b", ctx.text))
    stop = {"USA", "US", "U.S", "PDF", "DOCX", "CEO", "CFO", "FYI", "ETA", "HTML", "HTTP", "HTTPS"}
    acronyms = {a for a in acronyms if a not in stop}

    undefined: List[str] = []
    for ac in sorted(acronyms):
        if not _acronym_defined(ctx.text, ac):
            undefined.append(ac)

    metrics["acronyms_found"] = len(acronyms)
    metrics["acronyms_undefined"] = len(undefined)

    penalty = min(max_points, 1.0 * len(undefined[:5]))
    metrics["penalty_points"] = round(penalty, 2)
    metrics["scoring_notes"] = "1 point per undefined acronym (capped at 5)."
    score = max_points - penalty

    for ac in undefined[: ctx.options.max_issues_per_check]:
        match = re.search(rf"\b{re.escape(ac)}\b", ctx.text)
        if match:
            issues.append(
                issue_from_span(
                    ctx,
                    "acronyms",
                    f"Acronym '{ac}' appears without a nearby definition (e.g., 'Name ({ac})').",
                    match.start(),
                    match.end(),
                )
            )
        else:
            issues.append(Issue("acronyms", f"Acronym '{ac}' appears without a nearby definition (e.g., 'Name ({ac})')."))

    return CheckResult(
        check_id="acronyms",
        name="Acronym definitions",
        score=round(score, 2),
        max_score=max_points,
        issues=issues,
        metrics=metrics,
        status="ok",
    )


CHECKS: List[CheckSpec] = [
    CheckSpec("language_tool", "Grammar and style (LanguageTool)", 35.0, check_language_tool),
    CheckSpec("readability", "Readability", 10.0, check_readability),
    CheckSpec("punctuation", "Punctuation and style", 10.0, check_punctuation_style),
    CheckSpec("whitespace", "Whitespace and formatting", 5.0, check_whitespace_formatting),
    CheckSpec("percentages", "Numbers and percentages consistency", 5.0, check_percentage_consistency),
    CheckSpec("headings", "Headings and capitalization", 5.0, check_headings_capitalization),
    CheckSpec("passive_voice", "Passive voice (heuristic)", 5.0, check_passive_voice),
    CheckSpec("parallel_lists", "Parallel structure in lists", 10.0, check_parallel_bullets),
    CheckSpec("links", "Links and citation formatting", 5.0, check_links_formatting),
    CheckSpec("acronyms", "Acronym definitions", 10.0, check_acronym_definitions),
]


# -----------------------------
# Scoring aggregation
# -----------------------------

def aggregate_report(
    file_path: str,
    text: str,
    language: str,
    options: AuditOptions,
    enabled_checks: Sequence[CheckSpec],
    doc_info: Optional[Dict[str, Any]] = None,
) -> AuditReport:
    ctx = build_context(text, language, options)
    results: List[CheckResult] = []

    total_weight = sum(c.weight for c in enabled_checks)
    available_weight = 0.0
    total_score = 0.0

    for check in enabled_checks:
        result = check.func(ctx, check.weight)
        results.append(result)

        if result.status == "ok":
            available_weight += check.weight
            total_score += result.score
        elif options.strict:
            available_weight += check.weight

    if options.max_issues and options.max_issues > 0:
        remaining = options.max_issues
        for result in results:
            if remaining <= 0:
                if result.issues:
                    result.metrics["issues_truncated"] = len(result.issues)
                    result.issues = []
                continue
            if len(result.issues) > remaining:
                result.metrics["issues_truncated"] = len(result.issues) - remaining
                result.issues = result.issues[:remaining]
                remaining = 0
            else:
                remaining -= len(result.issues)

    coverage = 100.0 * available_weight / total_weight if total_weight else 0.0
    normalized_score = 0.0
    if available_weight > 0:
        normalized_score = round(100 * total_score / available_weight, 2)

    summary = {
        "file": file_path,
        "language": language,
        "score": normalized_score,
        "coverage": round(coverage, 2),
        "configured_max": total_weight,
        "scored_max": available_weight,
        "checks": {r.name: {"score": r.score, "max": r.max_score, "status": r.status} for r in results},
    }
    doc_info = doc_info or {}
    if doc_info:
        summary["document"] = doc_info

    return AuditReport(
        file_path=file_path,
        language=language,
        total_score=normalized_score,
        coverage=round(coverage, 2),
        check_results=results,
        summary=summary,
        doc_info=doc_info,
    )


def report_to_json(audit: AuditReport) -> Dict[str, Any]:
    return {
        "file_path": audit.file_path,
        "language": audit.language,
        "total_score": audit.total_score,
        "coverage": audit.coverage,
        "document": audit.doc_info,
        "checks": [
            {
                "id": c.check_id,
                "name": c.name,
                "score": c.score,
                "max_score": c.max_score,
                "points_lost": round(c.max_score - c.score, 2) if c.status == "ok" else None,
                "status": c.status,
                "metrics": c.metrics,
                "issues": [asdict(i) for i in c.issues],
            }
            for c in audit.check_results
        ],
    }


def report_to_markdown(audit: AuditReport) -> str:
    lines = []
    lines.append("# Document Audit Report")
    lines.append(f"**File:** `{audit.file_path}`  ")
    lines.append(f"**Language:** {audit.language}  ")
    lines.append(f"**Total Score:** **{audit.total_score}/100**  ")
    lines.append(f"**Coverage:** {audit.coverage}%")
    if audit.doc_info:
        lines.append("**Document Info:**")
        for key, value in audit.doc_info.items():
            lines.append(f"- {key}: {value}")
    lines.append("")
    for c in audit.check_results:
        lines.append(f"## {c.name} - {c.score:.2f}/{c.max_score:.0f} ({c.status})")
        points_lost = round(c.max_score - c.score, 2) if c.status == "ok" else None
        lines.append(f"**Points lost:** {points_lost if points_lost is not None else 'n/a'}")
        if c.metrics:
            lines.append("**Metrics**:")
            for k, v in c.metrics.items():
                lines.append(f"- {k}: {v}")
        if c.issues:
            lines.append("**Issues:**")
            for i, iss in enumerate(c.issues, 1):
                loc = f" ({iss.location})" if iss.location else ""
                lines.append(f"{i}. {iss.message}{loc}")
        lines.append("")
    return "\n".join(lines)


def _highlight_snippet(snippet: str, highlight: Optional[str]) -> str:
    if not snippet:
        return ""
    if not highlight:
        return html.escape(snippet)
    try:
        idx = snippet.find(highlight)
    except Exception:
        idx = -1
    if idx < 0:
        return html.escape(snippet)
    pre = html.escape(snippet[:idx])
    mid = html.escape(highlight)
    post = html.escape(snippet[idx + len(highlight):])
    return f"{pre}<mark>{mid}</mark>{post}"


def report_to_html(audit: AuditReport) -> str:
    lines = []
    lines.append("<!doctype html>")
    lines.append("<html lang=\"en\">")
    lines.append("<head>")
    lines.append("<meta charset=\"utf-8\"/>")
    lines.append("<meta name=\"viewport\" content=\"width=device-width, initial-scale=1\"/>")
    lines.append("<title>Document Audit Report</title>")
    lines.append("<style>")
    lines.append("body{font-family:Georgia,serif;margin:24px;color:#1b1b1b;background:#faf8f3;}")
    lines.append(".summary{padding:16px;background:#fff8e7;border:1px solid #f0e2be;border-radius:10px;}")
    lines.append(".check{margin-top:20px;padding:14px;background:#ffffff;border:1px solid #e8e0d0;border-radius:10px;}")
    lines.append(".check h2{margin:0 0 6px 0;font-size:18px;}")
    lines.append(".meta{color:#444;font-size:14px;margin:6px 0;}")
    lines.append(".issues{margin:8px 0;padding-left:18px;}")
    lines.append("code,pre{background:#f6f2ea;padding:6px;border-radius:6px;display:block;white-space:pre-wrap;}")
    lines.append("mark{background:#ffdf91;padding:0 2px;}")
    lines.append("</style>")
    lines.append("</head>")
    lines.append("<body>")
    lines.append("<h1>Document Audit Report</h1>")
    lines.append("<div class=\"summary\">")
    lines.append(f"<div><strong>File:</strong> {html.escape(audit.file_path)}</div>")
    lines.append(f"<div><strong>Language:</strong> {html.escape(audit.language)}</div>")
    lines.append(f"<div><strong>Total Score:</strong> {audit.total_score}/100</div>")
    lines.append(f"<div><strong>Coverage:</strong> {audit.coverage}%</div>")
    if audit.doc_info:
        lines.append("<div><strong>Document Info:</strong></div>")
        lines.append("<ul>")
        for key, value in audit.doc_info.items():
            lines.append(f"<li>{html.escape(str(key))}: {html.escape(str(value))}</li>")
        lines.append("</ul>")
    lines.append("</div>")

    for c in audit.check_results:
        points_lost = round(c.max_score - c.score, 2) if c.status == "ok" else None
        lines.append("<div class=\"check\">")
        lines.append(f"<h2>{html.escape(c.name)} ({html.escape(c.status)})</h2>")
        lines.append(f"<div class=\"meta\">Score: {c.score:.2f}/{c.max_score:.0f} | Points lost: {points_lost if points_lost is not None else 'n/a'}</div>")
        if c.metrics:
            lines.append("<div class=\"meta\"><strong>Metrics</strong></div>")
            lines.append("<ul>")
            for k, v in c.metrics.items():
                lines.append(f"<li>{html.escape(str(k))}: {html.escape(str(v))}</li>")
            lines.append("</ul>")
        if c.issues:
            lines.append("<div class=\"meta\"><strong>Issues</strong></div>")
            lines.append("<ol class=\"issues\">")
            for iss in c.issues:
                loc = f" ({iss.location})" if iss.location else ""
                lines.append(f"<li>{html.escape(iss.message)}{html.escape(loc)}</li>")
                if iss.snippet:
                    snippet_html = _highlight_snippet(iss.snippet, iss.highlight)
                    lines.append(f"<code>{snippet_html}</code>")
            lines.append("</ol>")
        lines.append("</div>")

    lines.append("</body>")
    lines.append("</html>")
    return "\n".join(lines)


def report_to_sarif(audit: AuditReport) -> Dict[str, Any]:
    rules: Dict[str, Any] = {}
    results: List[Dict[str, Any]] = []

    for check in audit.check_results:
        rule_id = check.check_id
        if rule_id not in rules:
            rules[rule_id] = {
                "id": rule_id,
                "name": check.name,
                "shortDescription": {"text": check.name},
            }
        for iss in check.issues:
            location: Dict[str, Any] = {
                "physicalLocation": {
                    "artifactLocation": {"uri": audit.file_path},
                }
            }
            region: Dict[str, Any] = {}
            if iss.line:
                region["startLine"] = iss.line
            if iss.col:
                region["startColumn"] = iss.col
            if iss.snippet:
                region["snippet"] = {"text": iss.snippet}
            if region:
                location["physicalLocation"]["region"] = region
            results.append(
                {
                    "ruleId": rule_id,
                    "level": "warning",
                    "message": {"text": iss.message},
                    "locations": [location],
                    "properties": {
                        "check_name": check.name,
                        "status": check.status,
                        "page": iss.page,
                    },
                }
            )

    sarif: Dict[str, Any] = {
        "version": "2.1.0",
        "$schema": "https://json.schemastore.org/sarif-2.1.0.json",
        "runs": [
            {
                "tool": {
                    "driver": {
                        "name": "doc_audit",
                        "rules": list(rules.values()),
                    }
                },
                "results": results,
            }
        ],
    }
    return sarif


# -----------------------------
# CLI
# -----------------------------

def _argv_value(argv: Sequence[str], flag: str) -> Optional[str]:
    for i, arg in enumerate(argv):
        if arg == flag and i + 1 < len(argv):
            return argv[i + 1]
        if arg.startswith(flag + "="):
            return arg.split("=", 1)[1]
    return None


def _argv_has_flag(argv: Sequence[str], flag: str) -> bool:
    return any(arg == flag or arg.startswith(flag + "=") for arg in argv)


def _listify(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(v) for v in value]
    if isinstance(value, str):
        return [value]
    return [str(value)]


def load_config(path: Optional[str]) -> Dict[str, Any]:
    if not path:
        return {}
    path = os.path.expanduser(path)
    if not os.path.exists(path):
        raise RuntimeError(f"Config file not found: {path}")
    ext = os.path.splitext(path)[1].lower()
    with open(path, "r", encoding="utf-8") as f:
        raw = f.read()
    if ext == ".toml":
        if tomllib is None:
            raise RuntimeError("TOML config requires Python 3.11+ or tomllib.")
        data = tomllib.loads(raw)
    elif ext in (".yaml", ".yml"):
        if yaml is None:
            raise RuntimeError("YAML config requires PyYAML. Install with: pip install pyyaml")
        data = yaml.safe_load(raw)
    else:
        data = json.loads(raw)
    if data is None:
        return {}
    if not isinstance(data, dict):
        raise RuntimeError("Config file must define a top-level object.")
    return data


def config_defaults(config: Dict[str, Any], argv: Sequence[str]) -> Dict[str, Any]:
    defaults: Dict[str, Any] = {}
    if "lang" in config:
        defaults["lang"] = config["lang"]
    if "out" in config:
        defaults["out"] = config["out"]
    if "max_issues" in config:
        defaults["max_issues"] = config["max_issues"]
    if "max_issues_per_check" in config:
        defaults["max_issues_per_check"] = config["max_issues_per_check"]
    if "lt_chunk_size" in config:
        defaults["lt_chunk_size"] = config["lt_chunk_size"]
    if "lt_overlap" in config:
        defaults["lt_overlap"] = config["lt_overlap"]
    if "strict" in config:
        defaults["strict"] = config["strict"]
    if "no_langtool" in config:
        defaults["no_langtool"] = config["no_langtool"]
    if "fail_under" in config:
        defaults["fail_under"] = config["fail_under"]
    if "min_coverage" in config:
        defaults["min_coverage"] = config["min_coverage"]

    outputs = config.get("outputs", {})
    if "html" in outputs:
        defaults["html"] = bool(outputs.get("html"))
    if "sarif" in outputs:
        defaults["sarif"] = bool(outputs.get("sarif"))

    ocr = config.get("ocr", {})
    if "mode" in ocr:
        defaults["ocr"] = ocr["mode"]
    if "lang" in ocr:
        defaults["ocr_lang"] = ocr["lang"]
    if "dpi" in ocr:
        defaults["ocr_dpi"] = ocr["dpi"]
    if "min_chars" in ocr:
        defaults["ocr_min_chars"] = ocr["min_chars"]

    if "enable" in config and not _argv_has_flag(argv, "--enable"):
        defaults["enable"] = _listify(config["enable"])
    if "disable" in config and not _argv_has_flag(argv, "--disable"):
        defaults["disable"] = _listify(config["disable"])
    return defaults

def parse_check_list(values: Sequence[str]) -> List[str]:
    items: List[str] = []
    for val in values:
        for part in val.split(","):
            part = part.strip()
            if part:
                items.append(part)
    return items


def resolve_out_prefix(file_path: str, out_arg: Optional[str], multiple_files: bool) -> str:
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    if not out_arg:
        return os.path.splitext(file_path)[0]
    out_arg = os.path.expanduser(out_arg)
    if out_arg.endswith(os.sep) or (os.path.exists(out_arg) and os.path.isdir(out_arg)):
        return os.path.join(out_arg, base_name)
    if multiple_files:
        return f"{out_arg}_{base_name}"
    return out_arg


def list_checks() -> None:
    print("Available checks:")
    for c in CHECKS:
        print(f"- {c.check_id}: {c.name} (weight {c.weight})")


def main() -> int:
    argv = sys.argv[1:]
    cfg_path = _argv_value(argv, "--config")
    config = load_config(cfg_path) if cfg_path else {}
    cfg_defaults = config_defaults(config, argv)
    config_files = _listify(config.get("files")) if "files" in config else []

    parser = argparse.ArgumentParser(description="Audit writing quality of DOCX/PDF/TXT/MD and score out of 100.")
    parser.add_argument(
        "files",
        nargs="*",
        default=[DEFAULT_FILE],
        help=f"Path(s) to input document(s) (default: {DEFAULT_FILE})",
    )
    parser.add_argument("--config", default=None, help="Path to JSON/TOML/YAML config file")
    parser.add_argument("--lang", default="en-US", help="Language code for grammar checks (default: en-US)")
    parser.add_argument("--out", default=None, help="Output path prefix or directory")
    parser.add_argument("--disable", action="append", default=[], help="Check ids to disable (comma-separated)")
    parser.add_argument("--enable", action="append", default=[], help="Only run these check ids (comma-separated)")
    parser.add_argument("--list-checks", action="store_true", help="List available checks and exit")
    parser.add_argument("--no-langtool", action="store_true", help="Skip LanguageTool checks")
    parser.add_argument("--langtool", action="store_false", dest="no_langtool", help="Force LanguageTool checks on")
    parser.add_argument("--max-issues", type=int, default=0, help="Max total issues to report (0 for unlimited)")
    parser.add_argument("--max-issues-per-check", type=int, default=50, help="Max issues per check")
    parser.add_argument("--lt-chunk-size", type=int, default=8000, help="LanguageTool chunk size")
    parser.add_argument("--lt-overlap", type=int, default=200, help="LanguageTool chunk overlap")
    parser.add_argument("--strict", action="store_true", help="Treat skipped checks as zero score")
    parser.add_argument("--no-strict", action="store_false", dest="strict", help="Do not penalize skipped checks")
    parser.add_argument("--fail-under", type=float, default=None, help="Exit non-zero if score is below this")
    parser.add_argument("--min-coverage", type=float, default=None, help="Exit non-zero if coverage is below this percent")
    parser.add_argument("--html", action="store_true", default=False, help="Write HTML report")
    parser.add_argument("--no-html", action="store_false", dest="html", help="Disable HTML report")
    parser.add_argument("--sarif", action="store_true", default=False, help="Write SARIF report")
    parser.add_argument("--no-sarif", action="store_false", dest="sarif", help="Disable SARIF report")
    parser.add_argument(
        "--ocr",
        choices=["off", "auto", "force"],
        default="off",
        help="OCR PDF input if needed (off|auto|force)",
    )
    parser.add_argument("--ocr-lang", default="eng", help="OCR language (tesseract language code)")
    parser.add_argument("--ocr-dpi", type=int, default=300, help="OCR DPI for PDF rendering")
    parser.add_argument("--ocr-min-chars", type=int, default=600, help="Auto-OCR threshold (min extracted chars)")
    parser.add_argument("-v", "--verbose", action="count", default=0, help="Increase logging verbosity")

    if cfg_defaults:
        parser.set_defaults(**cfg_defaults)

    args = parser.parse_args()

    if config_files and (args.files == [DEFAULT_FILE] or not args.files):
        args.files = config_files

    if args.list_checks:
        list_checks()
        return 0

    logging.basicConfig(
        level=logging.DEBUG if args.verbose and args.verbose > 1 else logging.INFO if args.verbose else logging.WARNING,
        format="%(levelname)s: %(message)s",
    )

    disable = set(parse_check_list(args.disable))
    enable = set(parse_check_list(args.enable))

    checks = CHECKS
    if enable:
        checks = [c for c in CHECKS if c.check_id in enable]
    if disable:
        checks = [c for c in checks if c.check_id not in disable]

    unknown = (enable | disable) - {c.check_id for c in CHECKS}
    if unknown:
        LOGGER.warning("Unknown check id(s): %s", ", ".join(sorted(unknown)))

    options = AuditOptions(
        max_issues=args.max_issues,
        max_issues_per_check=args.max_issues_per_check,
        lt_chunk_size=args.lt_chunk_size,
        lt_overlap=args.lt_overlap,
        strict=args.strict,
        enable_langtool=not args.no_langtool,
        ocr_mode=args.ocr,
        ocr_lang=args.ocr_lang,
        ocr_dpi=args.ocr_dpi,
        ocr_min_chars=args.ocr_min_chars,
    )

    exit_code = 0
    multiple_files = len(args.files) > 1

    for file_path in args.files:
        if not os.path.exists(file_path):
            LOGGER.error("File not found: %s", file_path)
            exit_code = max(exit_code, 2)
            continue

        try:
            doc = read_text_from_file(file_path, options)
            text = normalize_text(doc.text)
        except Exception as exc:
            LOGGER.error("Error reading file %s: %s", file_path, exc)
            exit_code = max(exit_code, 2)
            continue

        doc_info = {
            "pages": doc.page_count,
            "extraction_method": doc.extraction_method,
            "ocr_used": doc.ocr_used,
            "ocr_reason": doc.ocr_reason,
        }
        doc_info = {k: v for k, v in doc_info.items() if v is not None}

        audit = aggregate_report(file_path, text, args.lang, options, checks, doc_info=doc_info)

        out_prefix = resolve_out_prefix(file_path, args.out, multiple_files)
        out_dir = os.path.dirname(out_prefix)
        if out_dir:
            os.makedirs(out_dir, exist_ok=True)

        json_path = out_prefix + "_report.json"
        md_path = out_prefix + "_report.md"
        html_path = out_prefix + "_report.html"
        sarif_path = out_prefix + "_report.sarif"

        with open(json_path, "w", encoding="utf-8") as jf:
            json.dump(report_to_json(audit), jf, ensure_ascii=False, indent=2)

        with open(md_path, "w", encoding="utf-8") as mf:
            mf.write(report_to_markdown(audit))

        if args.html:
            with open(html_path, "w", encoding="utf-8") as hf:
                hf.write(report_to_html(audit))

        if args.sarif:
            with open(sarif_path, "w", encoding="utf-8") as sf:
                json.dump(report_to_sarif(audit), sf, ensure_ascii=False, indent=2)

        print(f"Score: {audit.total_score}/100 (coverage {audit.coverage}%) - {os.path.basename(file_path)}")
        for c in audit.check_results:
            print(f"  - {c.name}: {c.score:.2f}/{c.max_score:.0f} ({c.status})")
        print(f"\nSaved: {json_path}\nSaved: {md_path}")
        if args.html:
            print(f"Saved: {html_path}")
        if args.sarif:
            print(f"Saved: {sarif_path}")

        if args.fail_under is not None and audit.total_score < args.fail_under:
            exit_code = max(exit_code, 3)
        if args.min_coverage is not None and audit.coverage < args.min_coverage:
            exit_code = max(exit_code, 4)

    return exit_code


if __name__ == "__main__":
    sys.exit(main())
