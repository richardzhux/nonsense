from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_worksheet():
    doc = Document()

    # Define standard style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # --- PAGE 1: INTRODUCTION ---
    # Header Row
    header = doc.add_paragraph()
    header.add_run('Author: ').bold = True
    header.add_run('_________________________  ')
    header.add_run('Reviewer: ').bold = True
    header.add_run('_________________________')

    doc.add_heading('Introduction:', level=1)

    sections_pg1 = [
        ("Global-Level Concerns.", "(Does the introduction clearly differentiate the topic being studied from the argument being made about that topic? Does it signal who the thesis is in conversation with and how the thesis fills a gap in the literature? Does the author explain the topic in a way that assumes the reader’s basic familiarity with the field, or does it start too broadly or too narrowly?)"),
        ("Local-Level Concerns.", "(Does the introduction flow from topic to argument, general to specific? Does the author clearly differentiate what is debatable [i.e. their argument] from what is not [i.e. the facts about their topic]? Is the author appropriately selective about which secondary sources they mention in the introduction as being part of the academic conversation about their topic? Is the author narrating their personal process, or are they presenting their argument and findings in a scholarly voice? Do you have a sense of what kind of methodology the author used in this thesis?)"),
        ("Sentence-Level Concerns.", "(Does the author switch verb tenses? Are there typographical or spelling errors? Is a citation missing? Are there places where you would suggest an alternative word or phrase? If the author uses first person, are they consistent throughout and do they still retain a formal tone?)"),
        ("Based on the thesis paragraph – the paragraph that lays out the argument of the entire paper (probably located near or at the end of the introduction – how do you think the Analysis section is structured? List out the main points you expect it to hit and the order in which you think they appear.", "")
    ]

    for title, desc in sections_pg1:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        if desc:
            p.add_run(" " + desc)
        doc.add_paragraph("\n\n\n") # Space for writing

    footer1 = doc.add_paragraph("Please complete this side for the Introduction and the reverse side for the Conclusion.")
    footer1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- PAGE 2: CONCLUSION ---
    doc.add_heading('Conclusion:', level=1)

    sections_pg2 = [
        ("Global-Level Concerns.", "(Does the conclusion clearly belong to the same project set out in the Introduction? Does it assume that the reader has been paying attention throughout, or does it unnecessarily restate the main argument of the piece? Does the Conclusion engage with the work in a larger sociocultural context? Does the Conclusion “feel” like an ending?)"),
        ("Local-Level Concerns.", "(Does the introduction engage with the “so what”? Does it leave the reader thinking about either potential next steps for research or the larger implications of the project, or both? Does everything in the Conclusion clearly belong in the Conclusion, or are there points where the author continues trying to prove their main argument? If the Conclusion is longer than two or three paragraphs, where might the author narrow down their point?)"),
        ("Sentence-Level Concerns.", "(Does the author switch verb tenses? Are there typographical or spelling errors? Is a citation missing? Are there places where you would suggest an alternative word or phrase? If the author uses first person, are they consistent throughout? If the Conclusion takes more of an op-ed tone and thus integrates a bit of pathos and ethos, does it still stay in a formal register?)")
    ]

    for title, desc in sections_pg2:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(" " + desc)
        doc.add_paragraph("\n\n\n\n") # Extra space

    footer2 = doc.add_paragraph("Please complete this side for the Conclusion and the reverse side for the Introduction.")
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    file_name = "Thesis_Review_Worksheet.docx"
    doc.save(file_name)
    print(f"File saved as {file_name}")

if __name__ == "__main__":
    create_worksheet()